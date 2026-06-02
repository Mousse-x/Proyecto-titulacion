"""
document_processor.py
Servicio de procesamiento de documentos de evidencia LOTAIP.

Responsabilidades:
- Calcular hash SHA256 del archivo
- Detectar tipo MIME real
- Obtener metadatos del archivo
- Verificar si el archivo está corrupto o protegido
- Extraer texto del documento (PDF, XLSX, DOCX, CSV)
- Extraer tablas y columnas
- Guardar metadatos procesados en Evidence
"""

import csv
import hashlib
import os
import re
import logging
from pathlib import Path
from datetime import datetime
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from django.conf import settings
from django.utils import timezone

logger = logging.getLogger(__name__)

XLSX_NS = {
    "m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


# ──────────────────────────────────────────────────────────────────────
# HASH Y MIME
# ──────────────────────────────────────────────────────────────────────

def calculate_sha256(file_path):
    """Calcula el hash SHA256 de un archivo."""
    sha256 = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
    except Exception as e:
        logger.error(f"Error calculando SHA256: {e}")
        return None


def get_real_mime_type(file_path):
    """Detecta el tipo MIME real usando python-magic."""
    try:
        import magic
        mime = magic.Magic(mime=True)
        return mime.from_file(str(file_path))
    except Exception as e:
        logger.warning(f"Error detectando MIME con magic: {e}")
        # Fallback por extensión
        ext = Path(file_path).suffix.lower()
        mime_map = {
            ".pdf": "application/pdf",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".csv": "text/csv",
        }
        return mime_map.get(ext, "application/octet-stream")


def get_file_extension(file_path):
    """Obtiene la extensión real del archivo."""
    return Path(file_path).suffix.lower()


# ──────────────────────────────────────────────────────────────────────
# METADATOS
# ──────────────────────────────────────────────────────────────────────

def detect_lotaip_document_type(name_or_path):
    """Clasifica los 3 archivos LOTAIP esperados por literal."""
    text = Path(str(name_or_path or "")).name.lower()
    text = re.sub(r"[_\-.]+", " ", text)

    if "metadato" in text:
        return "Metadatos"
    if "diccionario" in text or "dicionario" in text:
        return "Diccionario"
    if "conjunto" in text and "dato" in text:
        return "Conjunto de datos"
    return None


def get_file_metadata(file_path):
    """Extrae metadatos básicos del archivo."""
    try:
        stat = os.stat(file_path)
        return {
            "file_name": Path(file_path).name,
            "extension": get_file_extension(file_path),
            "size_bytes": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }
    except Exception as e:
        logger.error(f"Error obteniendo metadatos: {e}")
        return {}


# ──────────────────────────────────────────────────────────────────────
# VERIFICACIÓN DE INTEGRIDAD
# ──────────────────────────────────────────────────────────────────────

def is_file_corrupted(file_path, mime_type=None):
    """Verifica si el archivo está corrupto intentando abrirlo."""
    ext = get_file_extension(file_path)

    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            doc.close()
            return False

        elif ext == ".xlsx":
            with ZipFile(file_path) as zf:
                return "xl/workbook.xml" not in zf.namelist()

        elif ext == ".xls":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(file_path), read_only=True)
                wb.close()
            except ImportError:
                logger.warning("openpyxl no disponible para verificar XLS; se asume accesible")
            return False

        elif ext == ".docx":
            from docx import Document
            Document(str(file_path))
            return False

        elif ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                f.read(1024)
            return False

        return False  # extensión desconocida, asumimos OK

    except Exception as e:
        logger.warning(f"Archivo posiblemente corrupto: {file_path} — {e}")
        return True


def is_password_protected(file_path):
    """Verifica si el archivo tiene contraseña/protección."""
    ext = get_file_extension(file_path)

    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(str(file_path))
            protected = doc.is_encrypted
            doc.close()
            return protected

        elif ext == ".xlsx":
            return False

        elif ext == ".xls":
            try:
                import openpyxl
                try:
                    wb = openpyxl.load_workbook(str(file_path), read_only=True)
                    wb.close()
                    return False
                except Exception:
                    return True
            except ImportError:
                return False

        elif ext == ".docx":
            from docx import Document
            try:
                Document(str(file_path))
                return False
            except Exception:
                return True

        return False
    except Exception:
        return False


# ──────────────────────────────────────────────────────────────────────
# EXTRACCIÓN DE TEXTO
# ──────────────────────────────────────────────────────────────────────

def extract_text_pdf(file_path):
    """Extrae texto de un archivo PDF usando pdfplumber."""
    text = ""
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                # Extraer tablas
                page_tables = page.extract_tables()
                for t in page_tables:
                    tables.append(t)
    except Exception as e:
        logger.error(f"Error extrayendo texto PDF: {e}")

    return text.strip(), tables


def _xlsx_shared_strings(zf):
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []

    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    return [
        "".join(t.text or "" for t in item.findall(".//m:t", XLSX_NS))
        for item in root.findall("m:si", XLSX_NS)
    ]


def _xlsx_cell_value(cell, shared_strings):
    value_node = cell.find("m:v", XLSX_NS)
    if value_node is None:
        return ""

    value = value_node.text or ""
    if cell.attrib.get("t") == "s" and value:
        try:
            return shared_strings[int(value)]
        except (IndexError, ValueError):
            return ""
    return value


def _extract_text_excel_xml(file_path):
    text = ""
    columns = []
    sheet_names = []
    row_count = 0

    with ZipFile(file_path) as zf:
        workbook = ET.fromstring(zf.read("xl/workbook.xml"))
        rels = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        rel_map = {rel.attrib["Id"]: rel.attrib["Target"] for rel in rels}
        shared_strings = _xlsx_shared_strings(zf)

        for sheet in workbook.findall("m:sheets/m:sheet", XLSX_NS):
            sheet_name = sheet.attrib.get("name", "")
            rel_id = sheet.attrib.get(f"{{{XLSX_NS['r']}}}id")
            target = rel_map.get(rel_id)
            if not target:
                continue

            sheet_names.append(sheet_name)
            sheet_path = f"xl/{target}" if not target.startswith("xl/") else target
            root = ET.fromstring(zf.read(sheet_path))
            rows = root.findall(".//m:sheetData/m:row", XLSX_NS)
            if not rows:
                continue

            parsed_rows = []
            for row in rows[:51]:
                values = [
                    _xlsx_cell_value(cell, shared_strings).strip()
                    for cell in row.findall("m:c", XLSX_NS)
                ]
                values = [value for value in values if value]
                if values:
                    parsed_rows.append(values)

            if parsed_rows:
                columns.extend(parsed_rows[0])
                for row in parsed_rows[:50]:
                    text += " | ".join(row) + "\n"

            row_count += max(0, len(rows) - 1)

    return text.strip(), list(dict.fromkeys(columns)), sheet_names, row_count


def extract_text_excel(file_path):
    """Extrae texto, columnas y datos de un archivo Excel."""
    text = ""
    columns = []
    sheet_names = []
    row_count = 0

    try:
        import pandas as pd

        ext = get_file_extension(file_path)
        engine = "openpyxl" if ext == ".xlsx" else None

        xls = pd.ExcelFile(str(file_path), engine=engine)
        sheet_names = xls.sheet_names

        for sheet in sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet, engine=engine)
            # Columnas
            sheet_cols = [str(c).strip() for c in df.columns.tolist() if str(c).strip() and not str(c).startswith("Unnamed")]
            columns.extend(sheet_cols)
            row_count += len(df)

            # Convertir contenido a texto
            for _, row in df.head(50).iterrows():  # Limitar a 50 filas para texto
                row_text = " | ".join([str(v) for v in row.values if pd.notna(v) and str(v).strip()])
                if row_text:
                    text += row_text + "\n"

        xls.close()

    except Exception as e:
        logger.warning(f"Extraccion Excel con pandas no disponible o fallo: {e}")
        if get_file_extension(file_path) == ".xlsx":
            try:
                return _extract_text_excel_xml(file_path)
            except Exception as fallback_error:
                logger.error(f"Error extrayendo datos Excel por XML: {fallback_error}")

    return text.strip(), list(set(columns)), sheet_names, row_count


def extract_text_docx(file_path):
    """Extrae texto y tablas de un archivo Word DOCX."""
    text = ""
    tables_data = []

    try:
        from docx import Document
        doc = Document(str(file_path))

        # Párrafos
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"

        # Tablas
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append(table_rows)
            # Agregar texto de tablas
            for row_data in table_rows:
                text += " | ".join(row_data) + "\n"

    except Exception as e:
        logger.error(f"Error extrayendo texto DOCX: {e}")

    return text.strip(), tables_data


def extract_text_csv(file_path):
    """Extrae texto y columnas de un archivo CSV."""
    text = ""
    columns = []
    row_count = 0

    try:
        for encoding in ["utf-8-sig", "utf-8", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding, newline="") as f:
                    sample = f.read(4096)
                    f.seek(0)
                    try:
                        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
                    except csv.Error:
                        dialect = csv.excel

                    reader = csv.reader(f, dialect)
                    rows = []
                    for idx, row in enumerate(reader):
                        rows.append([str(v).strip() for v in row])
                        if idx >= 100:
                            break

                if not rows:
                    break

                document_type = detect_lotaip_document_type(file_path)
                if document_type in {"Metadatos", "Diccionario"}:
                    columns = [row[0] for row in rows if row and row[0]]
                    data_rows = rows
                else:
                    columns = [c for c in rows[0] if c]
                    data_rows = rows[1:]
                row_count = len([r for r in data_rows if any(cell for cell in r)])

                for row in data_rows[:50]:
                    row_text = " | ".join([v for v in row if v])
                    if row_text:
                        text += row_text + "\n"
                break
            except UnicodeDecodeError:
                continue

    except Exception as e:
        logger.error(f"Error extrayendo texto CSV: {e}")

    return text.strip(), columns, row_count


# ──────────────────────────────────────────────────────────────────────
# ORQUESTADOR PRINCIPAL
# ──────────────────────────────────────────────────────────────────────

def process_document(evidence):
    """
    Procesa un documento de evidencia completo.
    Extrae metadatos, hash, texto, columnas, etc.
    Guarda los resultados en el modelo Evidence.

    Retorna un dict con toda la información extraída.
    """
    result = {
        "file_exists": False,
        "file_hash": None,
        "mime_type": None,
        "extension": None,
        "metadata": {},
        "is_corrupted": False,
        "is_protected": False,
        "extracted_text": "",
        "columns": [],
        "tables": [],
        "sheet_names": [],
        "row_count": 0,
        "document_type": None,
    }

    # Verificar existencia del archivo
    if not evidence.file_path:
        return result

    abs_path = Path(settings.MEDIA_ROOT) / evidence.file_path
    if not abs_path.exists():
        return result

    result["file_exists"] = True

    # Hash
    result["file_hash"] = calculate_sha256(abs_path)

    # MIME y extensión
    result["mime_type"] = get_real_mime_type(abs_path)
    result["extension"] = get_file_extension(abs_path)

    # Metadatos
    result["metadata"] = get_file_metadata(abs_path)
    result["document_type"] = detect_lotaip_document_type(f"{evidence.title} {abs_path.name}")

    # Integridad
    result["is_corrupted"] = is_file_corrupted(abs_path)
    result["is_protected"] = is_password_protected(abs_path)

    # Si está corrupto o protegido, no intentar extraer texto
    if result["is_corrupted"] or result["is_protected"]:
        _save_processing_results(evidence, result)
        return result

    # Extraer contenido según tipo
    ext = result["extension"]

    if ext == ".pdf":
        text, tables = extract_text_pdf(abs_path)
        result["extracted_text"] = text
        result["tables"] = tables
        # Intentar extraer columnas de las tablas PDF
        if tables:
            for t in tables:
                if t and len(t) > 0:
                    # La primera fila generalmente son encabezados
                    headers = [str(c).strip() for c in t[0] if c and str(c).strip()]
                    result["columns"].extend(headers)
            result["columns"] = list(set(result["columns"]))
            result["row_count"] = sum(len(t) - 1 for t in tables if len(t) > 1)

    elif ext in (".xlsx", ".xls"):
        text, columns, sheet_names, row_count = extract_text_excel(abs_path)
        result["extracted_text"] = text
        result["columns"] = columns
        result["sheet_names"] = sheet_names
        result["row_count"] = row_count

    elif ext == ".docx":
        text, tables = extract_text_docx(abs_path)
        result["extracted_text"] = text
        result["tables"] = tables

    elif ext == ".csv":
        text, columns, row_count = extract_text_csv(abs_path)
        result["extracted_text"] = text
        result["columns"] = columns
        result["row_count"] = row_count

    # Guardar resultados en el modelo
    _save_processing_results(evidence, result)

    return result


def _save_processing_results(evidence, result):
    """Guarda los resultados del procesamiento en la base de datos."""
    try:
        evidence.file_hash = result.get("file_hash")
        evidence.extracted_text = (result.get("extracted_text") or "")[:50000]  # Limitar a 50K chars
        evidence.metadata_json = {
            "mime_type": result.get("mime_type"),
            "extension": result.get("extension"),
            "metadata": result.get("metadata", {}),
            "is_corrupted": result.get("is_corrupted", False),
            "is_protected": result.get("is_protected", False),
            "columns": result.get("columns", []),
            "sheet_names": result.get("sheet_names", []),
            "row_count": result.get("row_count", 0),
            "document_type": result.get("document_type"),
        }
        evidence.processing_status = "processed" if result.get("file_exists") else "error"
        evidence.save(update_fields=["file_hash", "extracted_text", "metadata_json", "processing_status"])
    except Exception as e:
        logger.error(f"Error guardando resultados de procesamiento: {e}")
